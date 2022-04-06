from docx import Document
from docx.shared import Inches
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import json
from docx.shared import RGBColor
import os
import operator

def createProposal(d):
    f = open('data.json',)
    data = json.load(f)

    os.system('clear')

    code = input("Enter the difficulty seperated by spaces (Ex. 5 9 4 5): ").split(" ")
    today = date.today().strftime("%m-%d-%y")
    homeOwnerName = input("\nPlease input the Homeowner's Name(s): ")
    homeOwnerStreet = input("\nPlease input the Homeowner's Street: ")
    homeOwnerTown = input("\nPlease input the Homeowner's Town: ")
    homeOwnerZipcode = input("\nPlease input the Homeowner's Zipcode: ")
    homeOwnerPhone = input("\nPlease input the Homeowner's Phone Number (Ex. 845-742-5110): ")
    homeOwnerEmail = input("\nPlease `input the Homeowner's Email (Ex. ryan@rycorhvac.com): ")
    proposalAuthor = input("\nPlease enter YOUR full name (Ex. Scott Arnold): ")
    numOfUnits = int(input("\nHow many units?: "))

    units = {}
    totalPrice = 0
    totalBTU = 0
    for _ in range(numOfUnits):
        tempList = input("\nPlease enter location and type of unit (Ex. 18 Living Room): ").split(" ", 1)
        units[tempList[1]] = int(tempList[0])
        totalPrice = totalPrice + data["Price"][tempList[0]]
        totalBTU = totalBTU + data["BTU"][tempList[0]]

    print("\nTime to add discounts! If no more discounts, put 'n'. Format should be 'y' or 'n' if this is a per-unit discount, the amount, and reason.")
    totalDiscount = 0
    discount = 'y 200 Seasonal Discount'
    discounts = []
    discount = input("\nAdd discount (or enter 'n') (Ex. y 200 Seasonal Discount): ")
    while discount != 'n':
        discount = discount.split(" ", 2)
        discounts.append(
            {
                "per" : discount[0],
                "amount" : int(discount[1]),
                "reason" : discount[2]
            }
        )
        
        totalDiscount = totalDiscount + len(units) * int(discount[1]) if discount[0] == 'y' else totalDiscount + int(discount[1])

        discount = input("\nAdd discount (or enter 'n') (Ex. y 200 Seasonal Discount): ")

    print("\nTime to add notes! If no more notes, put 'n'. Format should just be whatever message you want!")
    notes = []
    note = input("\nAdd note (or enter 'n'): ")
    while note != 'n':
        notes.append(note)
        note = input("\nAdd note (or enter 'n'): ")

    units = {k: v for k, v in sorted(units.items(), key=lambda x: x[1], reverse=True)}

    t = '\t' * 9
    rycorString = t + 'RYCOR LLC\n' +t+ '135 North Chestnut Street\n' +t+ 'New Paltz, NY 12561\n' +t+ '845-742-5110\n' +t+ '%s\n' % (today) +t+ '%s,%s - %s,%s' % (code[0], code[1], code[2], code[3])

    homeOwnerString = '%s\n%s\n%s NY, %s\n%s\n%s' % (homeOwnerName, homeOwnerStreet, homeOwnerTown, homeOwnerTown, homeOwnerPhone, homeOwnerEmail)


    document = Document()

    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    document.styles['Normal'].font.name = "Times New Roman"
    document.styles['Normal'].font.size = Pt(12)







    document.add_picture('rycor.png')
    rycorHeader = document.add_paragraph(rycorString)
    rycorHeader.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    homeOwnerHeader = document.add_paragraph(homeOwnerString)

    note1Paragraph = document.add_paragraph()
    note1Paragraph.add_run('Client will sign Trade Ally Payment Authorization Form and utility company rebate will go directly to Rycor.').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    if notes:
        note2Paragraph = document.add_paragraph()
        for note in notes:
            note2Paragraph.add_run(note + '').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


    note3Paragraph = document.add_paragraph()
    note3Paragraph.add_run('Full Heat Solution.\n\n').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    note3Paragraph.add_run('ALL FS Models.\n\n').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    note3Paragraph.add_run('Mitsubishi Super High Efficiency Heating, Cooling, Dehumidification, and Air Purification system.')

    for loc, val in units.items():
        unitParagraph = document.add_paragraph('Install Mitsubishi ')
        unitParagraph.add_run("Hyper-Heat").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        unitParagraph.add_run(' Heat Pump in %s.' % (loc))
        provideString = f'Provide all materials and labor to install one {val},000 BTU- Mitsubishi Ductless Hyper-\tHeat Heat Pump system.'
        unitParagraph.add_run('\n\u25CF\t' + provideString)
        unitParagraph.add_run("\n\u25CF\tCondensing unit located on wall mount system.")
        unitParagraph.add_run("\n\u25CF\tEvaporator located close to ceiling.")
        unitParagraph.add_run("\n\u25CF\tAll lines will be covered in Line Hide Covering.")
        unitParagraph.add_run("\n\u25CF\tCondensation will be piped to outside.")
        unitPrice = data["Price"][str(val)]
        unitParagraph.add_run(f'\nTotal price to complete above mentioned work: ${unitPrice:,}')

    document.add_paragraph(f'Total price cost: ${totalPrice:,}')

    rebatePrice = int((totalBTU * 0.13)-500)

    rebateString = f'${rebatePrice:,} Central Hudson Rebate Mail-In Rebate. $1,300 off per 10,000 BTU for Full heating load installation. Client will sign Trade Ally Payment Authorization form and utility company rebate will go directly to Rycor'

    rebateAndPriceParagraph = document.add_paragraph(rebateString)

    discountString = ''
    for discount in discounts:
        amount = int(discount['amount'])
        reason = discount['reason']
        if discount['per'] == 'y':
            discountString = f'\n\n${(len(units) * amount):,} Discount- {reason}. ${amount} off per unit installed.'
        else:
            discountString = f'\n\n${amount} Discount- {reason}'
        
        rebateAndPriceParagraph.add_run(discountString).font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    rebateAndPriceParagraph.add_run(f'\n\nTotal due on day of installation: ${(totalPrice - rebatePrice - totalDiscount):,}').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    finalString = f'\n\nTotal Project Cost after rebates and discounts:\n${(totalPrice - rebatePrice - totalDiscount):,} Due on day of Installation via certified bank check'
    rebateAndPriceParagraph.add_run(finalString).font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    finalInformationParagraph = document.add_paragraph('12-year parts and compressor warrantee\n3-year labor warrantee')

    finalInformationParagraph.add_run("\n\nPermit's").bold = True

    permitString = ' - Permit and electrical inspection fees are not included and may vary by building department. RYCOR will provide all neccesary documents. Land owner will be responsible for bringing them to the building department.'
    finalInformationParagraph.add_run(permitString)

    finalInformationParagraph.add_run("\n\nNote:").bold = True

    electricString = ' Line Voltage and Panel work is not included. An electrician will be needed to complete this project. RYCOR recommends '
    finalInformationParagraph.add_run(electricString)
    finalInformationParagraph.add_run("Chris at Craft Electric 845-282-1514").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    document.add_paragraph('Thank you for taking the time to meet with me. It would be a pleasure to complete this project with you. Please feel free to call me at any time with any thoughts or questions.')

    t = '\t' * 6
    sincerelyString = t + 'Sincerely,\n' +t+ f'{proposalAuthor}\n' +t+ 'Comfort Specialist\n' +t+ '(845) 742-5110'
    sincerelyParagraph = document.add_paragraph(sincerelyString)
    sincerelyParagraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    unitList = ""
    for key, val in units.items():
        unitList += str(val) + " "

    document.save(f'{homeOwnerName} {unitList}Proposal.docx')