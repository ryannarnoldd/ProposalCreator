from docx import Document
from docx.shared import Inches
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import json
from docx.shared import RGBColor
import os
import operator

def createProposal(d):
    f = open('data.json',)
    data = json.load(f)

    #os.system('clear')

    today = date.today().strftime("%m-%d-%y")
    name = d['name']
    street = d['street']
    town = d['town']
    zipcode = d['zipcode']
    phone = d['phone']
    email = d['email']
    author = d['author']
    difficulty = d['difficulty'].split(" ")
    kumoCloud = d['kumoCloud']
    electric = d['electric']
    fullSolution = d['fullSolution']

    units = d['units']
    totalPrice = 0
    totalBTU = 0
    for loc in units:
        totalPrice = totalPrice + data["Price"][units[loc]]
        totalBTU = totalBTU + data["BTU"][units[loc]]

    units = {k: v for k, v in sorted(units.items(), key=lambda x: int(x[1]), reverse=True)}
     

    discounts = d['discounts']
    totalDiscount = 0
    for discount in discounts:
        if discount['perUnit']:
            totalDiscount = totalDiscount + len(units) * int(discount['amount'])
        else:
            totalDiscount + int(discount['amount'])

    notes = d['notes']

    t = '\t' * 9
    rycorString = t + 'RYCOR LLC\n' +t+ '135 North Chestnut Street\n' +t+ 'New Paltz, NY 12561\n' +t+ '845-742-5110\n' +t+ '%s\n' % (today) +t+ '%s,%s - %s,%s' % (difficulty[0], difficulty[1], difficulty[2], difficulty[3])

    homeOwnerString = '%s\n%s\n%s NY, %s\n%s\n%s' % (name, street, town, zipcode, phone, email)

    document = Document()

    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    document.styles['Normal'].font.name = "Times New Roman"
    document.styles['Normal'].font.size = Pt(12)







    document.add_picture('rycor.png')
    rycorHeader = document.add_paragraph(rycorString)
    rycorHeader.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    homeOwnerHeader = document.add_paragraph(homeOwnerString)

    noteParagraph = document.add_paragraph()
    noteParagraph.add_run('Client will sign Trade Ally Payment Authorization Form and utility company rebate will go directly to Rycor.').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    for note in notes:
            noteParagraph.add_run('\n\n' + note).font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


    noteParagraph.add_run('Full Heat Solution.\n\n').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    noteParagraph.add_run('ALL FS Models.\n\n').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    noteParagraph.add_run('Mitsubishi Super High Efficiency Heating, Cooling, Dehumidification, and Air Purification system.')

    for loc, val in units.items():
        unitParagraph = document.add_paragraph('Install Mitsubishi ')
        unitParagraph.add_run("Hyper-Heat").font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        unitParagraph.add_run(' Heat Pump in %s.' % (loc))
        provideString = f'Provide all materials and labor to install one {val},000 BTU- Mitsubishi Ductless Hyper-\tHeat Heat Pump system.'
        unitParagraph.add_run('\n\u25CF\t' + provideString)
        unitParagraph.add_run("\n\u25CF\tCondensing unit located on wall mount system.")
        unitParagraph.add_run("\n\u25CF\tEvaporator located close to ceiling.")
        unitParagraph.add_run("\n\u25CF\tAll lines will be covered in Line Hide Covering.")
        unitParagraph.add_run("\n\u25CF\tCondensation will be piped to outside.")
        unitPrice = data["Price"][str(val)]
        unitParagraph.add_run(f'\nTotal price to complete above mentioned work: ${unitPrice:,}')

    if kumoCloud:
        document.add_paragraph(f'${kumoCloud*295} Kumo Cloud - $295 for each Kumo Cloud device.')
        totalPrice += kumoCloud * 295

    document.add_paragraph(f'Total price cost: ${totalPrice:,}')
    if fullSolution:
        rebatePrice = int((totalBTU * 0.13)-500)
        rebateString = f'${rebatePrice:,} Central Hudson Rebate Mail-In Rebate. $1,300 off per 10,000 BTU for Full heating load installation. Client will sign Trade Ally Payment Authorization form and utility company rebate will go directly to Rycor'
    else:
        rebatePrice = 400 * len(units)
        rebateString = f'${rebatePrice:,} Rebate - Central Hudson Rebate Mail-In Rebate. $400 off per condenser for Partial heating load installation. Client will sign Trade Ally Payment Authorization form and utility company rebate will go directly to Rycor'


    rebateAndPriceParagraph = document.add_paragraph(rebateString)

    discountString = ''
    for discount in discounts:
        amount = int(discount['amount'])
        reason = discount['reason']
        if discount['perUnit']:
            discountString = f'\n\n${(len(units) * amount):,} Discount- {reason}. ${amount} off per unit installed.'
        else:
            discountString = f'\n\n${amount} Discount- {reason}'
        
        rebateAndPriceParagraph.add_run(discountString).font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    rebateAndPriceParagraph.add_run(f'\n\nTotal due on day of installation: ${(totalPrice - rebatePrice - totalDiscount):,}').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    finalString = f'\n\nTotal Project Cost after rebates and discounts:\n${(totalPrice - rebatePrice - totalDiscount):,} Due on day of Installation via certified bank check'
    rebateAndPriceParagraph.add_run(finalString).font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    finalInformationParagraph = document.add_paragraph('12-year parts and compressor warrantee\n3-year labor warrantee')

    finalInformationParagraph.add_run("\n\nPermit's").bold = True

    permitString = ' - Permit and electrical inspection fees are not included and may vary by building department. RYCOR will provide all neccesary documents. Land owner will be responsible for bringing them to the building department.'
    finalInformationParagraph.add_run(permitString)

    if electric != "Rycor HVAC":
        finalInformationParagraph.add_run("\n\nNote:").bold = True
        electricName = data['Electric'][electric]["Name"]
        electricPhone = data["Electric"][electric]["Phone"]
        electricString = ' Line Voltage and Panel work is not included. An electrician will be needed to complete this project. RYCOR recommends '
        finalInformationParagraph.add_run(electricString)
        finalInformationParagraph.add_run(f'{electricName} at {electric} {electricPhone}').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    document.add_paragraph('Thank you for taking the time to meet with me. It would be a pleasure to complete this project with you. Please feel free to call me at any time with any thoughts or questions.')

    t = '\t' * 6
    sincerelyString = t + 'Sincerely,\n' +t+ f'{author}\n' +t+ 'Comfort Specialist\n' +t+ '(845) 742-5110'
    sincerelyParagraph = document.add_paragraph(sincerelyString)
    sincerelyParagraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    unitList = ""
    for key, val in units.items():
        unitList += str(val) + " "

    filePath = f'/Users/ryanarnold/Desktop/Personal/Computer Science/Proposal/Proposals/{name} {street} {unitList}Proposal.docx'
    document.save(filePath)